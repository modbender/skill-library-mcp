#!/usr/bin/env python3
"""
DOCX 進階編輯器 - 支援表格、顯著文字、刪除線、格式保留
用法：
  uv run python scripts/docx_editor.py input.docx output.docx edits.json
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def apply_highlight(run):
    """應用螢光筆高亮"""
    run.font.highlight_color = True  # 黃色


def apply_strike(run):
    """應用刪除線"""
    run.font.strike = True


def apply_bold(run):
    """應用粗體"""
    run.font.bold = True


def apply_underline(run):
    """應用底線"""
    run.font.underline = WD_UNDERLINE.SINGLE


def apply_style(run, style_type):
    """對 run 應用樣式"""
    if style_type == "highlight":
        apply_highlight(run)
    elif style_type == "delete":
        apply_strike(run)
        apply_highlight(run)  # 刪除線 + 紅色背景
    elif style_type == "bold":
        apply_bold(run)
    elif style_type == "underline":
        apply_underline(run)


def replace_in_run(run, search_text, replace_text, style_type="replace"):
    """
    在 run 層級替換文字
    
    Returns: True if replaced, False otherwise
    """
    if search_text not in run.text:
        return False
    
    if style_type == "replace":
        run.text = run.text.replace(search_text, replace_text)
    else:
        # 分割文字並應用樣式
        parts = run.text.split(search_text)
        if len(parts) > 1:
            # 清除原 run 文字
            run.text = parts[0]
            # 新增替換文字（帶樣式）
            new_run = run.parent.add_run(replace_text)
            apply_style(new_run, style_type)
            # 新增剩餘文字
            if parts[1]:
                remaining_run = run.parent.add_run(parts[1])
                # 複製原格式
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
    
    return True


def replace_in_paragraph(para, search_text, replace_text, style_type="replace"):
    """在段落中替換文字，保留格式"""
    if search_text not in para.text:
        return False
    
    # 遍歷所有 run
    for run in list(para.runs):  # 使用 list() 避免在迭代時修改
        if replace_in_run(run, search_text, replace_text, style_type):
            return True
    
    return False


def replace_in_table(table, search_text, replace_text, style_type="replace"):
    """在表格中替換文字"""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, search_text, replace_text, style_type):
                    replaced = True
    return replaced


def add_text_after_paragraph(doc, after_text, add_text, style_type="highlight"):
    """在指定段落後新增文字"""
    for para in doc.paragraphs:
        if after_text in para.text:
            # 在段落後新增一個 run
            new_run = para.add_run(add_text)
            apply_style(new_run, style_type)
            return True
    return False


def edit_docx_advanced(input_path, output_path, edits):
    """
    進階編輯 DOCX 文件
    
    Args:
        input_path: 輸入文件
        output_path: 輸出文件
        edits: 編輯規則（JSON 格式）
    """
    doc = Document(input_path)
    stats = {"replacements": 0, "additions": 0, "deletions": 0}
    
    # 處理替換
    for edit in edits.get("replacements", []):
        search_text = edit["search"]
        replace_text = edit.get("replace", "")
        style = edit.get("style", "replace")
        
        # 替換段落
        for para in doc.paragraphs:
            if replace_in_paragraph(para, search_text, replace_text, style):
                stats["replacements"] += 1
        
        # 替換表格
        for table in doc.tables:
            if replace_in_table(table, search_text, replace_text, style):
                stats["replacements"] += 1
    
    # 處理新增
    for edit in edits.get("additions", []):
        after_text = edit["after"]
        add_text = edit["text"]
        style = edit.get("style", "highlight")
        
        # 在段落中新增
        for para in doc.paragraphs:
            if after_text in para.text:
                new_run = para.add_run(add_text)
                apply_style(new_run, style)
                stats["additions"] += 1
        
        # 在表格中新增
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if after_text in para.text:
                            new_run = para.add_run(add_text)
                            apply_style(new_run, style)
                            stats["additions"] += 1
    
    # 儲存
    doc.save(output_path)
    
    print(f"✅ 完成編輯：{output_path}")
    print(f"📊 統計：")
    print(f"   替換：{stats['replacements']} 處")
    print(f"   新增：{stats['additions']} 處")
    print(f"   刪除：{stats['deletions']} 處")
    
    return stats


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    edits_path = Path(sys.argv[3])
    
    if not input_path.exists():
        print(f"❌ 錯誤：文件不存在 {input_path}")
        sys.exit(1)
    
    if not edits_path.exists():
        print(f"❌ 錯誤：編輯規則文件不存在 {edits_path}")
        sys.exit(1)
    
    # 讀取編輯規則
    with open(edits_path, 'r', encoding='utf-8') as f:
        edits = json.load(f)
    
    # 執行編輯
    edit_docx_advanced(input_path, output_path, edits)


if __name__ == "__main__":
    main()
